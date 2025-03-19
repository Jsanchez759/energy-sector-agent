import requests
from bs4 import BeautifulSoup
import os
import json
from datetime import datetime
from utils.logger import logging, CustomException
import sys
import docx
import unicodedata
from llama_index.core import (
    Document,
    VectorStoreIndex,
    Settings
)
from llama_index.embeddings.ollama import OllamaEmbedding
from llama_index.llms.ollama import Ollama
from llama_index.vector_stores.chroma import ChromaVectorStore
import chromadb


class CREG:
    def __init__(self, url: str = "https://creg.gov.co/loader.php?lServicio=Documentos&lFuncion=infoCategoriaConsumo&tipo=RE"):
        self.url = url
        self.data_path = "src/creg/data"
        self.headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36"
        }
        self.month_index_translation = {
            "DIC": "DEC",
            "NOV": "NOV",
            "OCT": "OCT",
            "SEP": "SEP",
            "AGO": "AUG",
            "JUL": "JUL",
            "JUN": "JUN",
            "MAY": "MAY",
            "ABR": "APR",
            "MAR": "MAR",
            "FEB": "FEB",
            "ENE": "JAN",
        }
        self.embedding_model = OllamaEmbedding(
            model_name="mxbai-embed-large",
            base_url="http://localhost:11434",
            ollama_additional_kwargs={"mirostat": 0},
        )
        Settings.embed_model = self.embedding_model
        self.llm = Ollama(model="llama3.2:3b", request_timeout=120.0)
        Settings.llm = self.llm

    def detect_file_links(self):
        """ Detects file links from CREG website and returns 10 last resolutions

        Returns:
            documents (list): List of dictionaries containing the title and url of the document

        args:
            None
        """
        try:
            response = requests.get(self.url, headers=self.headers, timeout=10)
            soup = BeautifulSoup(response.text, "html.parser")
            links = soup.find_all("a", href=True)
            documents = [
                {"title": link.text.strip(), "url": link["href"]}
                for link in links
                if "ControlAdmin/BajarArchivo".lower() in link["href"].lower()
            ]
            return documents
        except Exception as e:
            logging.error(f"Error detecting file links: {CustomException(e, sys)}")
            return []

    def download_documents(self, documents: list) -> bool:
        """
        Downloads the documents from the CREG website

        Args:
            documents (list): List of dictionaries containing the title and url of the document

        Returns:
            sucess: True if the documents were downloaded successfully, False otherwise
        """
        try:
            for document in documents:
                filename = f"{document['url'][-12:]}.docx"
                filename = filename.replace("\\", "")
                url = document["url"]
                file_path = os.path.join(self.data_path, filename)
                response = requests.get(url, headers=self.headers, timeout=10)
                with open(file_path, "wb") as file:
                    file.write(response.content)

            logging.info(f"Downloaded latest 10 resolutions from CREG")
            return True
        except Exception as e:
            logging.error(f"Error downloading documents: {CustomException(e, sys)}")
            return False

    def remove_accents(self, text:str) -> str:
        """
        Removes accents from the text

        Args:
            text (str): Text to remove accents from

        Returns:
            text (str): Text without accents
        """
        return "".join(c for c in unicodedata.normalize("NFKD", text) if not unicodedata.combining(c))

    def process_documents(self, documents: list) -> bool:
        """
        Processes the documents downloaded from the CREG website and get text and metadata

        Args:
            documents (list): List of dictionaries containing the title and url of the document

        Returns:
            sucess: True if the documents were processed successfully, False otherwise
        """
        resolutions = []
        for resolution in documents:
            try:
                doc = docx.Document(f"{self.data_path}/{resolution}")
                resolution_metadata = {}
                resolution_values_para = [0,1]

                resolution_name = doc.paragraphs[resolution_values_para[0]].text
                while not resolution_name.startswith("RESOLUCIÓN"):
                    resolution_values_para[0] += 1
                    resolution_name = doc.paragraphs[resolution_values_para[0]].text

                resolution_name = self.remove_accents(resolution_name)

                resolution_date = doc.paragraphs[resolution_values_para[1]].text
                while not resolution_date.startswith("("):
                    resolution_values_para[1] += 1
                    resolution_date = doc.paragraphs[resolution_values_para[1]].text

                resolution_concept = doc.paragraphs[resolution_values_para[1]+1].text
                while resolution_concept == "":
                    resolution_values_para[1] += 1
                    resolution_concept = doc.paragraphs[resolution_values_para[1]+1].text

                resolution_concept = self.remove_accents(resolution_concept)

                resolution_date = resolution_date.replace("(", "").replace(")", "").replace(".", " ")
                day, month, year = resolution_date.split()
                month_en = self.month_index_translation.get(month, month)
                date_english = f"{day} {month_en} {year}"
                resolution_date = datetime.strptime(date_english, "%d %b %Y")

                full_text = [p.text for p in doc.paragraphs]
                full_text = [p for p in full_text if p]  # Remove empty paragraphs
                full_text = [self.remove_accents(p) for p in full_text]
                resolution_metadata = {
                    "name": resolution_name,
                    "resolution_date": resolution_date.strftime("%Y-%m-%d"),
                    "concept": resolution_concept,
                    "full_text": "\n".join(full_text),
                    "process_date": datetime.now().strftime("%Y-%m-%d"),
                }
                resolutions.append(resolution_metadata)
                # delete file
                os.remove(f"{self.data_path}/{resolution}")
                logging.info(f"Processed resolution: {resolution_metadata['name']}")    

            except Exception as e:
                logging.error(f"Error processing documents: {CustomException(e, sys)}")
                os.rename(f"{self.data_path}/{resolution}", f"{self.data_path}/to_check/{resolution}")
                continue

        if resolutions:
            file_name = f"resolutions_processed.json"
            folder = f"{self.data_path}/processed"
            file_path = os.path.join(folder, file_name)
            with open(file_path, "w") as file:
                json.dump(resolutions, file, indent=4, ensure_ascii=False)
            return True
        else:
            logging.error(f"No resolutions were processed")
            return False

    def model_resolution_doc(self):
        """
        Model the resolution document

        Args:
            None

        Returns:
            documents (list): List of documents modeled with llama index
        """
        try:
            file_path = f"{self.data_path}/processed/resolutions_processed.json"
            with open(file_path, "r", encoding="utf-8") as f:
                resolutions = json.load(f)
            documents = [
                Document(
                    text=resolution["full_text"],
                    metadata={
                        "name": resolution["name"],
                        "date": resolution["resolution_date"],
                        "concept": resolution["concept"],
                    },
                )
                for resolution in resolutions
            ]
            return documents
        except Exception as e:
            logging.error(f"Error modeling resolution document: {CustomException(e, sys)}")
            return []

    def get_creg_vector_store(self, documents: list) -> VectorStoreIndex:
        """
        Get the vector store for the CREG model

        Args:
            documents (list): List of documents modeled with llama index

        Returns:
            index (VectorStoreIndex): Vector store index for the CREG model
        """
        try:
            db2 = chromadb.PersistentClient(path="chroma_db")
            chroma_collection = db2.get_or_create_collection("creg_index")
            vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
            index = VectorStoreIndex.from_vector_store(
                vector_store,
                embed_model=self.embedding_model,
            )
            return index
        except Exception as e:
            logging.error(f"Error getting CREG vector store: {CustomException(e, sys)}")
            return None

    def get_query_engine(self, index: VectorStoreIndex):
        """
        Get the query engine for the CREG model

        Args:
            index (VectorStoreIndex): Vector store index for the CREG model

        Returns:
            query_engine (QueryEngine): Query engine for the CREG model
        """
        try:
            query_engine = index.as_query_engine()
            return query_engine
        except Exception as e:
            logging.error(f"Error getting query engine: {CustomException(e, sys)}")
            return None
